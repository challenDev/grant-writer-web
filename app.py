import streamlit as st
from datetime import datetime
from pathlib import Path
from docx import Document
import requests
import json
import io

# Page config
st.set_page_config(
    page_title="Grant Writer AI – Win More Funding",
    page_icon="📄",
    layout="centered"
)

# Load prompt
with open("prompts/grants_template.txt", "r", encoding="utf-8") as f:
    PROMPT_TEMPLATE = f.read()

# Initialize Groq API settings (using requests instead of OpenAI SDK for Python 3.14 compatibility)
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"
MODEL = "llama-3.3-70b-versatile"  # Latest & best as of Dec 2025

st.title("501(c)(3) Grant Proposal Writer AI")
st.markdown("*Professional, funder-ready proposals in under 60 seconds*")

st.info("Built for nonprofits · Uses Llama 3.3 70B via Groq · 100% free to use")

with st.expander("Instructions", expanded=False):
    st.write("""
    1. Fill out as much as you know (all fields are optional but more = better proposal)
    2. Click **Generate Proposal**
    3. Download your Word document and submit!
    """)

with st.form("grant_form"):
    col1, col2 = st.columns(2)
    with col1:
        funder_name = st.text_input("Funder Name*", placeholder="The Smith Family Foundation")
        program_name = st.text_input("Program/RFP", value="General Operating Support")
        amount_requested = st.text_input("Amount Requested*", placeholder="$75,000")
        deadline = st.text_input("Deadline", placeholder="March 15, 2026")
    with col2:
        project_name = st.text_input("Project Name*", placeholder="Youth STEM Afterschool Program")
        project_dates = st.text_input("Project Dates", placeholder="July 1, 2025 – June 30, 2026")
        org_name = st.text_input("Organization Name*", placeholder="Bright Futures Nonprofit")
        ein = st.text_input("EIN", placeholder="XX-XXXXXXX")

    mission = st.text_area("Mission Statement (1-2 sentences)*", height=80)
    need_statement = st.text_area("Problem/Need Statement – Why this matters now*", height=120)
    project_description = st.text_area("Full Project Description*", height=120)
    
    col3, col4 = st.columns(2)
    with col3:
        goals_objectives = st.text_area("Specific, Measurable Goals & Objectives*", height=100)
        timeline = st.text_area("Timeline*", height=100)
    with col4:
        budget_narrative = st.text_area("Budget Narrative (justify major expenses)*", height=100)
        evaluation_plan = st.text_area("Evaluation Plan – How will you measure success?*", height=100)

    sustainability = st.text_area("Sustainability Plan (after funding ends)*", height=100)
    funder_priorities = st.text_input("Funder’s Top Priorities", placeholder="racial equity, youth development, STEM access")
    organizational_capacity = st.text_area("Past Success & Organizational Capacity (brag!)", height=100)

    submitted = st.form_submit_button("🚀 Generate Professional Grant Proposal")

if submitted:
    if not all([funder_name, amount_requested, project_name, org_name, mission, need_statement, project_description]):
        st.error("Please fill all fields marked with *")
    else:
        if not GROQ_API_KEY or GROQ_API_KEY == "your-groq-api-key-here":
            st.error("❌ Groq API key not configured. Please add your key to `.streamlit/secrets.toml`")
        else:
            with st.spinner("Writing your winning proposal... (usually 20–40 seconds)"):
                data = {
                    "funder_name": funder_name,
                    "program_name": program_name or "General Support",
                    "amount_requested": amount_requested,
                    "deadline": deadline or "Rolling",
                    "project_name": project_name,
                    "project_dates": project_dates or "TBD",
                    "org_name": org_name,
                    "ein": ein or "Available upon request",
                    "year_founded": "",
                    "service_area": "",
                    "populations": "",
                    "annual_budget": "",
                    "executive_director": "",
                    "need_statement": need_statement,
                    "project_description": project_description,
                    "goals_objectives": goals_objectives,
                    "timeline": timeline,
                    "budget_narrative": budget_narrative,
                    "evaluation_plan": evaluation_plan,
                    "sustainability": sustainability,
                    "funder_priorities": funder_priorities or "mission alignment and impact",
                    "organizational_capacity": organizational_capacity or "Strong track record of impact"
                }

                prompt = PROMPT_TEMPLATE.format(**data)

                try:
                    # Call Groq API via requests
                    headers = {
                        "Authorization": f"Bearer {GROQ_API_KEY}",
                        "Content-Type": "application/json"
                    }
                    payload = {
                        "model": MODEL,
                        "messages": [{"role": "user", "content": prompt}],
                        "temperature": 0.3,
                        "max_tokens": 12000
                    }
                    
                    response = requests.post(GROQ_API_URL, json=payload, headers=headers, timeout=60)
                    response.raise_for_status()
                    
                    result = response.json()
                    proposal_text = result["choices"][0]["message"]["content"]

                    # Generate Word doc
                    doc = Document()
                    doc.add_heading(f"Grant Proposal: {project_name}", 0)
                    doc.add_paragraph(f"Generated on {datetime.now():%B %d, %Y}")
                    doc.add_page_break()

                    for para in proposal_text.split("\n\n"):
                        text = para.strip()
                        if text.startswith("# "):
                            doc.add_heading(text[2:], level=1)
                        elif text.startswith("## "):
                            doc.add_heading(text[3:], level=2)
                        elif text:
                            doc.add_paragraph(text)

                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    safe_filename = f"{funder_name} - {project_name}.docx".replace(" ", "_").replace("/", "-")

                    st.success("Proposal generated successfully!")
                    st.download_button(
                        label="📄 Download Word Document",
                        data=buffer.getvalue(),
                        file_name=safe_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    with st.expander("Preview (text version)", expanded=False):
                        st.markdown(proposal_text)
                        
                except requests.exceptions.RequestException as e:
                    st.error(f"❌ API Error: {str(e)}")
                    st.info("Please check your Groq API key and try again.")

st.markdown("---")
st.caption("Built with ❤️ using Groq + Llama 3.3 70B · Free for all 501(c)(3)s")